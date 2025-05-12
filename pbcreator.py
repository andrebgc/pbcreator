import csv
import random
import string
import os
import configparser
import codecs
import win32com.client as win32
from docxtpl import DocxTemplate, Listing
from docx2pdf import convert
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

##### CREATE BANNER

print("""
██████╗░██████╗░░█████╗░██████╗░███████╗░█████╗░████████╗░█████╗░██████╗░
██╔══██╗██╔══██╗██╔══██╗██╔══██╗██╔════╝██╔══██╗╚══██╔══╝██╔══██╗██╔══██╗
██████╔╝██████╦╝██║░░╚═╝██████╔╝█████╗░░███████║░░░██║░░░██║░░██║██████╔╝
██╔═══╝░██╔══██╗██║░░██╗██╔══██╗██╔══╝░░██╔══██║░░░██║░░░██║░░██║██╔══██╗
██║░░░░░██████╦╝╚█████╔╝██║░░██║███████╗██║░░██║░░░██║░░░╚█████╔╝██║░░██║
╚═╝░░░░░╚═════╝░░╚════╝░╚═╝░░╚═╝╚══════╝╚═╝░░╚═╝░░░╚═╝░░░░╚════╝░╚═╝░░╚═╝
""")
print("made by: André Coutinho                          version 1.2")
print("")

##### CREATE MENU

menu_options = {
    1: 'CLIENT1',
    2: 'CLIENT2',
    3: 'Exit'
}

def print_menu():
    for key in menu_options.keys():
        print (key, '--', menu_options[key] )

def option1():
    print('Generating Playbook docs to client: \'CLIENT1\'')
    global client_id
    client_id = "CL-01"

def option2():
     print('Generating Playbook docs to client: \'CLIENT2\'')
     global client_id
     client_id = "CL-02"

if __name__=='__main__':
    while(True):
        print_menu()
        option = ''
        try:
            option = int(input('Enter your choice: '))
        except:
            print('Wrong input. Please enter a number ...')
        #Check what choice was entered and act accordingly
        if option == 1:
            option1()
            break
        elif option == 2:
            option2()
            break
        elif option == 10:
            print('Bye!!')
            exit()
else:
    print('Invalid option. Please enter a number between 1 and 10.')

class DocumentTemplate:
    file = ""  # Template file location
    name = ""  # Name of template
    naming = ""  # Naming template for resulting files

# Configuration File Variable
conf = os.path.abspath(os.getcwd()) + "\\" + client_id + "\\" + f"{client_id}.conf"

main_config = configparser.RawConfigParser()
main_config.read_file(codecs.open(f"{conf}", "r", "utf8"))

# Variables to work with Excel file
raw_var = main_config["current"]["var_to_indicate_row"]

cdsn_var = main_config["current"]["complete_data_sheet_name"]

# Excel file that we use as Database
xlsm_input = os.path.abspath(os.getcwd()) + "\\" + "use_cases_db.xlsx"

templates = []  # list of templates

templates_coutn = int(main_config["current"]["docx_templates_count"])
for i in range(1, templates_coutn + 1):
    dt = DocumentTemplate()
    dt.name = main_config["current"]["docx_template_name_" + str(i)]
    dt.file = os.path.abspath(os.getcwd()) + "\\" + client_id + "\\" + "playbook-template.docx"
    dt.file2 = os.path.abspath(os.getcwd()) + "\\" + client_id + "\\" + "info-playbook-template.docx"
    dt.naming = main_config["current"]["docx_template_naming_" + str(i)]
    templates.append(dt)

save_dir = os.path.abspath(os.getcwd()) + "\\" + client_id + "\\" + "output"

current_dir = os.path.abspath(os.getcwd())

img_dir = os.path.join(current_dir, f"{cdsn_var}")

excel = win32.Dispatch("Excel.Application")

excel.Visible = False
excel.DisplayAlerts = False

wb = excel.Workbooks.Open(xlsm_input)

random_name = ''.join(random.choices(string.ascii_lowercase + string.digits, k=4))
data_file = os.environ["temp"] + "\\" + random_name + ".csv"

wb.Sheets(main_config["current"]["complete_data_sheet_name"]).Select()

wb.SaveAs(data_file, 23)

wb.Close()
excel.Quit()
try:
    with open(data_file, "r") as csv_file:
        dialect = csv.Sniffer().sniff(csv_file.readline())
        csv_file.seek(0)
        csv_reader = csv.reader(csv_file, dialect)

        rows = list()
        for row in csv_reader:
            rows.append(row)

        doc_vars = rows[0]
        raw_var_index = 0
        for i in range(0, len(doc_vars)):
            if doc_vars[i] == raw_var:
                raw_var_index = i
                break

        first_data_row = 1 + int(main_config["current"]["lines_to_skip"])
        for i in range(first_data_row, len(rows)):
            if str(rows[i][raw_var_index]) == "":
                continue
            data = dict()
            for l in range(0, len(rows[0])):
                if doc_vars[l] == '':
                    continue
                if "\n" in str(rows[i][l]):
                    data[str(doc_vars[l])] = Listing(str(rows[i][l]))  # Listing is used to send multiline strings
                else:
                    data[str(doc_vars[l])] = str(rows[i][l])

            for t in templates:
                doc = DocxTemplate(t.file)
                doc2 = DocxTemplate(t.file2)
                doc.render(data)
                doc2.render(data)                                 
                Table = doc.tables[1]
                Table2 = doc.tables[3]
                Table3 = doc2.tables[1]
                Table4 = doc2.tables[3]
                #GET CELLS XML ELEMENT
                cell_xml_element = Table.rows[0].cells[1]._tc
                #RETRIEVE THE TABLE CELL PROPERTIES
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                #CREATE SHADING OBJECT
                shade_obj = OxmlElement('w:shd')
                color_table = Table.rows[0].cells[1].text
                if color_table == "Alta" or color_table == "Alto" or color_table == "High":
                    #SET THE SHADING OBJECT With RED
                    shade_obj.set(qn('w:fill'), "FF0000")
                    #APPEND THE PROPERTIES TO THE TABLE CELL PROPERTIES
                    table_cell_properties.append(shade_obj)
                elif color_table == "Media" or color_table == "Média" or color_table == "Medio" or color_table == "Médio" or color_table == "Medium":
                    #SET THE SHADING OBJECT With ORANGE
                    shade_obj.set(qn('w:fill'), "FFA500")
                    #APPEND THE PROPERTIES TO THE TABLE CELL PROPERTIES
                    table_cell_properties.append(shade_obj)
                elif color_table == "Baixa" or color_table == "Baixo" or color_table == "Low":
                    #SET THE SHADING OBJECT With YELLOW
                    shade_obj.set(qn('w:fill'), "FFFF00")
                    #APPEND THE PROPERTIES TO THE TABLE CELL PROPERTIES
                    table_cell_properties.append(shade_obj)                    
                else:
                    #SET THE SHADING OBJECT With GRAY
                    shade_obj.set(qn('w:fill'), "808080")
                    #APPEND THE PROPERTIES TO THE TABLE CELL PROPERTIES
                    table_cell_properties.append(shade_obj)
                                                        
                info = {"Informacional", "Informativo", "Info", "Informational"}
                # Creating de Playbooks documents in docx and pdf
                if str(rows[i][raw_var_index]) == "X" and color_table not in info:
                    # INSERT WORKFLOW FILE IMAGE    
                    p = Table2.rows[0].cells[0].add_paragraph()
                    r = p.add_run()
                    r.add_picture(img_dir + "/" + f"{cdsn_var}.jpeg" ,width=Inches(7.0), height=Inches(1.5))
                    # CREATE FILE AND CONVERT PDF
                    s = f"PLY-{t.naming}" % data
                    print(f"Generating Analytical Playbook - " , s.replace('.docx', ''))
                    doc.save(save_dir + "/" + s)
                    convert(save_dir + "/" + s)
                else:
                    # INSERT WORKFLOW FILE IMAGE
                    p = Table4.rows[0].cells[0].add_paragraph()
                    r = p.add_run()
                    r.add_picture(img_dir + "/" + f"{cdsn_var}.jpeg" ,width=Inches(7.0), height=Inches(1.5))
                    # CREATE FILE AND CONVERT PDF
                    s = f"PLY-{t.naming}" % data
                    print(f"Generating Informational Playbook - " , s.replace('.docx', ''))
                    doc2.save(save_dir + "/" + s)
                    convert(save_dir + "/" + s)
                    
except Exception as err:
    print(err)
finally:
    if os.path.exists(data_file):
        os.remove(data_file)
print("PB Creator Done!!!!!!!")
